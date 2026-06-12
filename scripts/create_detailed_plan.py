from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/kalien/Desktop/HerbWorld_Share_项目详细策划书_2026-06-12.docx")
LOGO = Path("/Users/kalien/Downloads/herbworldshare_logo_transparent.png")

GREEN = "173229"
DARK = "102019"
GOLD = "B28A35"
LIGHT_GREEN = "E9F0EB"
LIGHT_GOLD = "F5F0E3"
LIGHT_GRAY = "F2F4F3"
MUTED = "667168"
RED = "8B3A33"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def font(run, size=10.5, bold=False, color=DARK, name="PingFang SC"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.28

    for name, size, color, before, after in [
        ("Title", 28, GREEN, 0, 10),
        ("Subtitle", 13, MUTED, 0, 16),
        ("Heading 1", 17, GREEN, 18, 9),
        ("Heading 2", 13.5, GOLD, 13, 6),
        ("Heading 3", 11.5, GREEN, 9, 4),
    ]:
        style = styles[name]
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.22


def add_p(doc, text="", bold_lead=None, align=None, color=DARK, size=10.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), size=size, bold=True, color=color)
        font(p.add_run(text[len(bold_lead):]), size=size, color=color)
    else:
        font(p.add_run(text), size=size, color=color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        font(p.add_run(item))


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        font(p.add_run(item))


def add_callout(doc, title, text, fill=LIGHT_GOLD, title_color=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width(cell, 9120)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, 160, 180, 160, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run(title), 11, True, title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    font(p2.add_run(text), 10.5, False, DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths=None, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    if widths is None:
        widths = [9120 // len(headers)] * len(headers)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, value in enumerate(headers):
        cell = header.cells[i]
        set_cell_width(cell, widths[i])
        set_cell_fill(cell, GREEN)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(str(value)), font_size, True, "FFFFFF")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx % 2:
                set_cell_fill(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or (isinstance(value, str) and value.startswith("$")) else WD_ALIGN_PARAGRAPH.LEFT
            font(p.add_run(str(value)), font_size, False, DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    font(run, 9, False, MUTED)


def configure_section(section, first_page=False):
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.34)
    section.footer_distance = Inches(0.34)
    section.different_first_page_header_footer = first_page
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font(hp.add_run("HERBWORLD SHARE  |  PROJECT EXECUTION PLAN"), 8.5, True, MUTED)
    fp = section.footer.paragraphs[0]
    font(fp.add_run("Polaris Global L.L.C.  ·  机密内部执行文件  ·  "), 8.5, False, MUTED)
    add_page_number(fp)


def cover(doc):
    section = doc.sections[0]
    configure_section(section, True)
    for _ in range(1):
        doc.add_paragraph()
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture = p.add_run().add_picture(str(LOGO), width=Inches(1.15))
        doc_pr = picture._inline.docPr
        doc_pr.set("descr", "HerbWorld Share 品牌标志")
        doc_pr.set("title", "HerbWorld Share")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    font(p.add_run("HERBWORLD SHARE"), 12, True, GOLD)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("项目详细策划书"), 30, True, GREEN)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("从曼哈顿真实物品再利用试点，到可信城市互助网络"), 14, False, MUTED)
    doc.add_paragraph()
    add_callout(
        doc,
        "核心决策",
        "未来 90 天不以下载量、帖子数或复杂功能为目标，只验证三件事：能否稳定完成真实交接、能否控制安全与人工成本、能否从自然产生的本地服务和机构需求中获得收入。",
        LIGHT_GREEN,
        GREEN,
    )
    for _ in range(1):
        doc.add_paragraph()
    add_p(doc, "运营主体：Polaris Global L.L.C.", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED)
    add_p(doc, "首个运营区域：美国纽约市曼哈顿", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED)
    add_p(doc, "版本日期：2026 年 6 月 12 日", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED)
    doc.add_page_break()


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def build():
    doc = Document()
    setup_styles(doc)
    cover(doc)

    h1(doc, "一、执行摘要与项目判断")
    add_p(doc, "HerbWorld Share 是一个以附近免费物品互助为入口、以经过一次性取货码验证的真实交接为核心资产、以本地服务与机构再利用项目为近期收入来源的平台。它解决的不是“哪里可以发免费物品”这一单点问题，而是把发布、审核、匹配、隐私保护、现场交接、结果核验和后续服务组织成一个可信闭环。")
    add_p(doc, "项目已经具备可运营的早期基础：网站已上线，曼哈顿试点页面、物品发布、领取申请、服务需求、后台审核、举报、安全规则、隐私与条款页面、一次性取货码核销、Neon PostgreSQL 接口和 SwiftUI iOS 原型均已存在。当前最重要的工作不是继续扩展概念，而是把已有系统用于真实试点，并补齐账户、权限、审计、图片处理和运营流程。")
    add_callout(doc, "一句话定位", "让送出一件仍有用的物品，比把它扔掉更容易；让每一次真实交接都可验证、可复盘、可形成城市减废与服务需求数据。")
    add_table(doc, ["判断项", "当前结论", "经营含义"], [
        ["产品需求", "免费闲置流转需求已被市场验证，但单纯社区平台直接收费困难", "免费互助负责获客和信任，收入必须来自自然衍生服务与机构项目"],
        ["差异化", "一次性取货码、人工安全审核、最少地址披露和可审计结果构成差异", "不能把“发布量”当资产，必须把“真实完成交接”当资产"],
        ["启动区域", "曼哈顿范围仍偏大，首批运营应聚焦一个高密度街区或合作物业", "先形成局部供需密度，再扩大地理范围"],
        ["收入时点", "首批收入不应来自用户会员费或交易抽成", "先做商家成交推荐费、物业批量项目费和协调服务费"],
        ["最大风险", "安全事故、地址泄露、失约、虚假交接和人工成本失控", "增长前必须建立审核、权限、证据、举报与事故响应闭环"],
    ], [1400, 3600, 4120])

    h2(doc, "未来 12 个月的目标")
    add_bullets(doc, [
        "完成至少 100 次经取货码核销的真实交接，并把完成率、失约率、事故率和每次交接人工分钟数记录清楚。",
        "与至少 3 家经过主体、许可与保险核验的本地服务商完成真实付费订单。",
        "与至少 1 个物业、学校、办公室或社区机构完成批量再利用付费试点。",
        "完成账户、角色权限、最少地址披露、申诉、数据删除和运营审计能力。",
        "在核心闭环稳定前，不上线用户间付款、资金托管、保险、可交易积分或复杂游戏。",
    ])

    h1(doc, "二、项目现状审计")
    h2(doc, "已经完成的资产")
    add_table(doc, ["资产", "当前状态", "可立即用于什么"], [
        ["品牌与主体", "HerbWorld Share；运营主体 Polaris Global L.L.C.；网站 herbworld.app", "统一对外身份、商标与 App Store 组织申请"],
        ["商标申请", "Section 1(b) 意向使用基础、两个类别，$700 已支付；仍需确认申请序列号", "建立商标状态台账，等待审查并准备未来使用证据"],
        ["网站产品", "发布、浏览、领取申请、服务需求、举报、安全、条款、隐私、合作商家页面", "立即开展人工运营试点"],
        ["运营后台", "物品四项审核、申请处理、取货码生成、服务需求和举报处理", "支撑低量人工审核与真实闭环"],
        ["验证机制", "六位一次性取货码、72 小时有效、哈希保存、核销后完成", "把点击与真实交接区分开"],
        ["技术基础", "Next.js、Neon PostgreSQL、Vercel、API 限流、风险关键词、Telegram 通知", "支撑试点和快速迭代"],
        ["iOS 原型", "SwiftUI 工程、隐私清单、App Store 元数据和截图", "完成账户与 API 接入后进入 TestFlight"],
        ["运营资料", "商业计划、产品原则、30 天手册、安全研究与运营 SOP", "直接用于团队培训和伙伴沟通"],
    ], [1500, 3600, 4020])

    h2(doc, "上线前必须补齐的缺口")
    add_table(doc, ["优先级", "缺口", "完成标准", "建议期限"], [
        ["P0", "确认 USPTO 申请序列号与正式收据", "在 Trademark Center/TSDR 可查到申请，保存脱敏收据与序列号", "立即"],
        ["P0", "后台身份与权限", "运营后台不只依赖单一共享密钥；启用 MFA、角色和审计日志", "14 天"],
        ["P0", "账户与隐私生命周期", "注册、账户删除、数据导出、同意版本、保留期限均可执行", "30 天"],
        ["P0", "安全事件台账", "每次举报、下架、申诉、恢复和事故都有证据与负责人", "7 天"],
        ["P1", "图片上传与隐私处理", "对象存储、EXIF 清除、敏感信息检查、生命周期删除", "30 天"],
        ["P1", "申请选择与候补流程", "发布者可选择领取者；取消、失约、候补状态完整", "30 天"],
        ["P1", "商家核验台账", "营业主体、保险、适用许可、价格、投诉与复核日期齐全", "21 天"],
        ["P2", "附近发现与通知", "距离带、保存搜索、推送通知；不公开精确地址", "90 天后按数据决定"],
    ], [900, 2600, 4100, 1520])

    h1(doc, "三、市场、用户与竞争策略")
    h2(doc, "核心用户与高频触发场景")
    add_table(doc, ["用户群", "真正要完成的任务", "最强触发时刻", "平台切入方式"], [
        ["发布者", "快速处理仍可使用的物品并腾出空间", "搬家、换租、装修、整理房屋", "低摩擦发布、人工协调、必要时提供搬运/清走报价"],
        ["领取者", "在附近获得真正需要的物品", "刚搬家、学生开学、新移民安家", "附近发现、申请说明、守约记录和安全交接"],
        ["物业/机构", "减少遗留物品处置成本并保留结果记录", "搬出季、宿舍清退、办公室搬迁", "批量录入、预约、匹配和影响报告"],
        ["本地商家", "获得信息完整、可成交的本地客户", "工作日空档、低峰时段、特定路线", "预审需求、限定区域、成交后收费"],
    ], [1500, 2650, 2200, 2770])
    h2(doc, "竞争格局与差异化")
    add_p(doc, "Buy Nothing、Freecycle、Facebook 群组、Craigslist、Nextdoor 和 OLIO 已经证明免费流转需求真实存在。HerbWorld Share 不应试图靠“也能发免费物品”竞争，而应把差异集中在可信交接、运营安全、结果数据和自然衍生服务。")
    add_table(doc, ["能力", "普通群组/分类信息", "HerbWorld Share 目标能力"], [
        ["公开隐私", "常见直接公开联系方式或通过私聊交换地址", "公开仅显示模糊区域，确认后才最少开放必要信息"],
        ["内容安全", "依赖群管理员或事后举报", "发布前人工审核、召回核验、禁止品规则、事故台账"],
        ["结果真实性", "点击、留言或标记已取走即可", "一次性取货码现场核销后才记录完成"],
        ["服务衔接", "用户自行寻找搬运、清洁、组装", "在明确场景下连接经过核验的本地商家"],
        ["机构价值", "很难提供完整结果与影响记录", "提供批量项目管理和可解释的影响报告"],
    ], [1800, 3400, 3920])

    h1(doc, "四、产品设计与用户全流程")
    h2(doc, "产品分层")
    add_table(doc, ["层级", "作用", "首发原则"], [
        ["实用层", "发布、浏览、申请、确认和举报", "必须永久免费且几秒内能理解"],
        ["信任层", "审核、最少地址披露、取货码、举报、申诉和审计", "信任是核心产品，不是客服附属功能"],
        ["影响层", "真实再利用记录与可解释的减废估算", "只展示有依据的估算，不夸大环保结果"],
        ["运动层", "可选 Apple Health、步数转非现金游戏能量", "不阻挡核心功能，不出售或训练健康数据"],
        ["世界层", "真实交接点亮城市与文化世界", "只在真实互助闭环稳定后投入"],
    ], [1300, 3650, 4170])
    h2(doc, "完整用户流程")
    add_steps(doc, [
        "发布者选择“我有东西”，提交物品、状态、缺陷、尺寸、楼层、电梯、可交接时间和模糊区域。",
        "系统先检查禁止物品、联系方式、完整门牌、重复内容与异常提交；运营人员完成所有权、状态、召回和隐私四项审核。",
        "审核通过后公开。公开页面不展示完整地址、私人联系方式或精确坐标。",
        "领取者提交可领取时间、搬运计划和必要说明；发布者或运营人员选择领取者，并保留候补。",
        "双方确认交接窗口后，系统生成限时一次性取货码，仅向领取者提供。",
        "现场检查无误后，发布者核销取货码；系统将交接记为真实完成，并结束敏感信息访问。",
        "交接前后在明确场景询问是否需要搬运、清洁、组装或合规清走服务；由用户与经过核验的商家直接签约付款。",
        "发生危险、骚扰、收费、状态严重不符或召回问题时，立即停止交接、下架内容、保留证据并进入举报与申诉流程。",
    ])
    h2(doc, "可信等级与公平匹配")
    add_table(doc, ["等级", "条件", "权限与限制"], [
        ["访客", "无需注册", "浏览公开物品和安全规则；不能申请或查看交接信息"],
        ["基础真人", "Apple/Google/邮箱注册并验证；接受规则", "少量发布和申请；异常行为触发人工复核"],
        ["可信成员", "至少一次真实核销，历史无严重违规", "更高同时申请额度、更快审核和更强通知"],
        ["增强验证", "高价值、高频或异常行为时触发", "设备、号码、动态物品验证或第三方身份核验；必须明确同意"],
    ], [1400, 3500, 4220])
    add_callout(doc, "公平原则", "贡献记录可以提高匹配机会，但不得让新用户或弱势用户永久失去参与资格。建议匹配权重：需求匹配 30%、距离与取货能力 25%、守约与安全 20%、真实贡献 15%、新用户与随机公平机会 10%。", LIGHT_GREEN, GREEN)

    h1(doc, "五、商业模式与单位经济")
    h2(doc, "收入顺序")
    add_table(doc, ["阶段", "收入来源", "收费方式", "何时启动"], [
        ["近期", "经过核验的本地服务商合作", "完成订单后的固定推荐费或 10%–15% 营销服务费", "30 天试点即可验证"],
        ["近期", "物业与机构批量再利用项目", "按项目、地点或活动日收费", "完成 20–50 次真实交接后销售"],
        ["中期", "商家运营工具与需求订阅", "按月订阅，含区域和需求通知", "有稳定成交量后"],
        ["中期", "物业订阅与影响报告", "按地点/月或年度合同", "机构试点续约后"],
        ["长期", "可选高级便利功能", "Apple IAP 或网页订阅；核心互助保持免费", "账户与留存稳定后"],
        ["长期", "合规聚合数据产品", "合同授权与用途限定", "建立同意账本、脱敏和审计后"],
    ], [1100, 2750, 3270, 2000])
    h2(doc, "服务订单单位经济示例")
    add_table(doc, ["项目", "示例金额", "说明"], [
        ["商家公开价", "$200", "用户自行寻找商家的参考价"],
        ["平台用户价", "$185", "商家用低峰产能和较低获客成本支持专属价"],
        ["商家支付平台", "$20", "仅在订单完成后支付"],
        ["商家净收", "$165", "必须高于商家可接受的边际收益"],
        ["用户获得价值", "$15 + 信息预审与协调", "不能只靠折扣，必须减少沟通与报价摩擦"],
        ["平台毛收入", "$20", "扣除人工协调、退款争议与获客费用后才是贡献毛利"],
    ], [2600, 1700, 4820])
    add_p(doc, "关键管理指标不是订单总价，而是每个服务需求的有效率、报价率、成交率、每单协调分钟数、投诉率和平台贡献毛利。若单笔推荐费长期无法覆盖人工协调成本，应提高预审标准、改用项目费，或停止该服务品类。")

    h1(doc, "六、冷启动与曼哈顿试点")
    h2(doc, "30 天验证方案")
    add_table(doc, ["周期", "动作", "最低目标", "停止或调整条件"], [
        ["第 1 周", "访谈搬运、清洁、组装各 5 家；选定唯一试点微区域", "3 家提供书面平台价和成交后费用", "没有 3 家愿意合作，先调整商家价值主张"],
        ["第 2 周", "从现有客户、经纪、物业、搬家公司和学生群收集首批物品", "15 件合格物品、5 次真实交接、2 个服务需求", "发布多但无人申请，缩小品类或调整区域"],
        ["第 3 周", "优化审核、提醒、候补和报价预审", "累计 30 件发布、10 次交接、1 笔服务订单", "人工时间过高，暂停扩量并修流程"],
        ["第 4 周", "回访所有参与者与商家，计算单位经济", "15 次交接、5 个服务需求、2 笔订单", "任何关键指标未达标，延长试点而非开发复杂 App"],
    ], [1100, 3500, 2500, 2020])
    h2(doc, "首批供给获取")
    add_bullets(doc, [
        "房产经纪和租赁客户：把搬出前可用物品集中发布，帮助客户降低处理摩擦。",
        "公寓物业：从一栋楼或一位物业经理开始，建立固定搬出日流程。",
        "搬家公司、整理师与清洁公司：他们最早看到客户准备丢弃的可用物品。",
        "留学生与新移民群体：一端产生集中供给，另一端产生明确需求。",
        "创始人及真实朋友网络：用于跑通第一批安全监督交接，不制造虚假数据。",
    ])
    h2(doc, "90 天里程碑")
    add_table(doc, ["阶段", "核心结果", "产品交付", "经营决策"], [
        ["0–30 天", "15 次真实交接、2 笔服务订单", "修复最关键流程缺口，建立事故与商家台账", "判断本地服务需求是否自然发生"],
        ["31–60 天", "累计 50 次真实交接、至少 1 个物业试点", "账户、候补、提醒、图片隐私处理", "计算每次交接与每单服务的人力成本"],
        ["61–90 天", "累计 100 次真实交接、形成重复参与", "距离带、基础通知、运营数据看板", "决定继续曼哈顿扩张、换区域或暂停"],
    ], [1300, 2850, 3100, 1870])

    h1(doc, "七、运营、安全与反作弊")
    h2(doc, "每日运营最低动作")
    add_bullets(doc, [
        "每天至少两次检查新发布、领取申请、服务需求和举报；高风险举报优先处理。",
        "批准发布前完成所有权、状态、CPSC 召回和隐私四项审核，并记录审核人和时间。",
        "只向已核验主体、保险和适用许可的商家发送服务需求；报价必须有范围、总价、取消和损坏责任。",
        "完整地址只在双方确认且确有必要时开放；任务结束后关闭访问并按保留规则删除。",
        "每周复盘真实交接、失约、举报、事故、每次交接人工分钟数和每单贡献毛利。",
    ])
    h2(doc, "反作弊控制")
    add_table(doc, ["风险信号", "自动控制", "人工处理"], [
        ["短期大量领取同类或高价值物品", "限制同时申请数、延迟贡献入账", "核对用途、历史和关联账号"],
        ["重复图片、网络图片、批量账号", "图片哈希、设备/IP/号码关系提示", "动态物品验证或拒绝发布"],
        ["账号群反复互相核销", "图谱异常告警、同物品仅一个有效码", "冻结贡献、保留证据、申诉复核"],
        ["高频失约或引导线下收费", "降低额度、限制功能", "警告、暂停或封禁关联账号"],
        ["危险、召回或状态严重不符", "立即下架和冻结流程", "事故台账、通知相关用户并评估报告义务"],
    ], [2850, 3000, 3270])
    h2(doc, "关键服务等级")
    add_table(doc, ["指标", "试点目标", "升级警戒线"], [
        ["普通物品审核中位时间", "12 小时内", "连续两周超过 24 小时"],
        ["高风险举报首次响应", "2 小时内", "任何严重举报超过 4 小时"],
        ["申请到确认中位时间", "24 小时内", "超过 48 小时且无提醒"],
        ["真实交接完成率", "≥ 60%", "低于 40%"],
        ["失约率", "≤ 15%", "高于 25%"],
        ["每次交接人工运营时间", "≤ 25 分钟", "高于 45 分钟"],
        ["严重安全事件率", "目标为 0", "任何事件立即暂停相关流程复盘"],
    ], [3300, 2800, 3020])

    h1(doc, "八、技术架构、数据安全与产品路线")
    h2(doc, "当前架构与建议")
    add_table(doc, ["领域", "当前基础", "下一步要求"], [
        ["客户端", "Next.js 网站 + SwiftUI iOS 原型", "优先保证网站真实运营；iOS 完成账户与 API 后进入 TestFlight"],
        ["服务端", "Next.js Route Handlers", "增加正式身份、角色权限、审计、限流和任务队列"],
        ["数据库", "Neon PostgreSQL", "公开数据与敏感数据分离；备份、恢复演练和最小权限"],
        ["媒体", "当前以 URL/后续联系收图为主", "对象存储、EXIF 清除、病毒检查、敏感信息检查和生命周期"],
        ["通知", "Telegram 运营通知", "用户端使用邮件/APNs；敏感内容不进入通用通知正文"],
        ["风控", "关键词、联系方式/地址检测、请求限流", "设备信号、重复图片、账号关系、申诉和证据留存"],
    ], [1700, 3300, 4120])
    h2(doc, "数据治理规则")
    add_bullets(doc, [
        "每个字段必须有明确用途、访问角色、保留期限、删除行为和是否可用于统计的说明。",
        "联系方式、完整地址、私人消息、身份证件和 HealthKit 数据不得进入外部数据产品或 AI 训练数据。",
        "公开位置只使用模糊街区或距离带；精确地址单独加密并按任务临时授权。",
        "后台敏感操作必须记录操作者、时间、资源和理由；高权限账户启用 MFA 和硬件安全密钥。",
        "建立数据泄露响应：隔离、保全日志、评估影响、依法通知、修复控制并复盘。",
    ])
    h2(doc, "12 个月产品路线")
    add_table(doc, ["时间", "必须交付", "暂不做"], [
        ["0–30 天", "账户最小版、候补/取消/失约、事故台账、商家核验表、真实运营看板", "复杂地图、支付、游戏"],
        ["31–90 天", "图片存储与隐私处理、通知、距离带、数据删除和导出、角色审计", "全国开放、自动批准"],
        ["4–6 月", "贡献记录、异常领取风控、机构批量工具、中英西关键流程", "可交易积分、现金奖励"],
        ["7–12 月", "第二城市门槛验证、影响地图、可选运动与世界层小范围原型", "在核心指标不稳时扩展完整游戏"],
    ], [1200, 4700, 2420])

    h1(doc, "九、公司、法律、保险与商标后续")
    add_callout(doc, "法律边界", "本策划书是经营与执行文件，不构成法律意见。新增付款、托管、保险、捐款、现金池、配送劳动者、用户间交易或跨州扩张前，应进行专项法律审查。", "F8E8E6", RED)
    h2(doc, "当前低风险运营边界")
    add_bullets(doc, [
        "平台不拥有、储存、专业检验或运输用户物品；物品必须完全免费。",
        "平台不处理用户间付款、押金或资金托管，不承诺保险或服务担保。",
        "付费服务由经过核验的独立商家向用户书面报价，用户与商家直接签约付款。",
        "危险品、食品、药品、安全关键产品、召回品和无法确认所有权的物品不开放。",
        "纽约州内搬家业务核验适用 NYSDOT 权限；跨州搬家核验 FMCSA 权限。",
    ])
    h2(doc, "公司与常规合规费用")
    add_table(doc, ["事项", "官方/市场费用", "时间与说明"], [
        ["纽约外国 LLC Application for Authority", "$250 官方申请费", "是否必须申请应由纽约律师/CPA 按实际经营活动判断"],
        ["纽约外国 LLC 公告", "$500–$2,500+ 预算范围", "获授权后 120 天内，在指定日报和周报连续六周；实际价格取决于县和报纸"],
        ["纽约 Biennial Statement", "$9 / 每两年", "获授权后按申请月份每两年申报"],
        ["Apple Developer Program", "$99 / 年", "组织账户需要 D-U-N-S、域名邮箱和公开有效网站"],
        ["注册代理、会计、税务申报", "$1,500–$6,000 / 年预算", "按主体、州别、收入与申报复杂度询价"],
        ["商业综合责任与网络保险", "$2,000–$10,000 / 年预算", "上线前由经纪人按业务边界、收入和风险正式报价"],
        ["隐私、条款、商家合同法律审查", "$5,000–$20,000 初期预算", "支付、劳务、数据产品等功能另做专项审查"],
    ], [2800, 2450, 3870])
    h2(doc, "HerbWorld Share 商标费用台账")
    add_table(doc, ["事项", "两个类别金额", "状态/触发条件"], [
        ["Section 1(b) 基础申请费", "$700", "已支付；需确认申请序列号与正式提交状态"],
        ["Statement of Use / Amendment to Allege Use", "$300", "实际在商业中使用商标并提交合格证据时，$150/类"],
        ["Statement of Use 延期", "$250 / 每次", "尚未具备使用证据时，$125/类；每次通常延长六个月"],
        ["若错过期限申请恢复", "$250", "仅在符合条件时；应以避免错过期限为主"],
        ["第 5–6 年 Section 8 使用声明", "$650", "$325/类，不含律师费"],
        ["可选 Section 8 + 15", "$1,150", "$575/类，用于使用声明与不可争议声明组合"],
        ["第 10 年 Section 8 + 9 续展", "$1,300", "$650/类，之后每十年一次"],
        ["商标律师与 Office Action", "$1,500–$6,000+ 预留", "取决于审查意见、复杂度和律师报价"],
    ], [3100, 2200, 3820])
    add_p(doc, "管理动作：建立商标日历，记录申请序列号、审查状态、Office Action 截止日、Notice of Allowance 日期、首次商业使用日期、使用证据和每次付款收据。已支付的 $700 不重复计入未来预算；未来至少预留 $300 的使用声明官方费，并为延期或审查意见准备应急预算。")

    h1(doc, "十、团队与治理")
    add_table(doc, ["角色", "试点期配置", "主要责任", "何时扩编"], [
        ["创始人/总负责人", "1 人，全职或核心投入", "产品、运营、合作、重大风险、现金管理", "始终保留最终责任"],
        ["技术负责人", "创始人兼任或合同制", "网站、App、数据库、安全、发布", "每月交接 >300 或出现持续技术瓶颈"],
        ["社区运营与审核", "兼职 1 人起", "审核、协调、举报、回访、台账", "审核 SLA 连续两周不达标"],
        ["安全与客服", "早期轮值/兼职", "高风险事件、申诉、夜间升级", "严重事件或服务时段扩大"],
        ["伙伴与机构销售", "创始人主导", "商家核验、物业与机构项目", "付费项目可覆盖岗位成本"],
        ["律师、CPA、保险经纪", "外部专业顾问", "公司、隐私、合同、税务和保险", "按功能与城市专项启用"],
    ], [1700, 2300, 3050, 2070])
    add_callout(doc, "治理原则", "每一项新功能或新城市必须写明：解决的问题、验证指标、最高月成本、合规负责人、停止条件。没有负责人和停止条件的功能不进入开发。", LIGHT_GREEN, GREEN)

    h1(doc, "十一、三年费用预算与收入预测")
    add_p(doc, "以下为基准经营模型，使用 2026 年美元，目的是确定现金需求和阶段门槛，不是对收入的保证。模型包含合理的合同开发、运营和合规成本，但不包含外部融资成本；实际执行应按月滚动更新。")
    h2(doc, "关键假设")
    add_bullets(doc, [
        "第一年聚焦曼哈顿微区域，完成 1,000 次真实交接；第二年扩展至纽约多个区域并验证第二城市；第三年形成多城市运营。",
        "居民核心发布和领取免费；收入来自服务推荐、物业/机构项目、商家工具与后期可选高级便利功能。",
        "服务推荐平均平台收入按每单 $25–$35 估算；机构项目按规模和运营复杂度单独报价。",
        "创始人薪资可根据现金情况延后，但模型必须承认真实团队成本，避免低估长期资金需求。",
    ])
    h2(doc, "三年基准预算")
    budget_rows = [
        ["人员与合同开发", "$24,000", "$120,000", "$330,000", "产品、工程、设计与维护"],
        ["社区运营、审核与客服", "$24,000", "$90,000", "$240,000", "随交接量和服务时段增长"],
        ["法律、会计、税务与保险", "$22,000", "$35,000", "$70,000", "含纽约与新增城市专项审查"],
        ["云、短信、媒体、监控与软件", "$6,000", "$24,000", "$90,000", "设置硬上限与媒体生命周期"],
        ["市场、社区活动与伙伴开发", "$12,000", "$45,000", "$120,000", "以局部密度和机构合作为主"],
        ["硬件、安全密钥与办公", "$4,000", "$8,000", "$20,000", "早期不购买物理服务器"],
        ["应急与不可预见支出", "$8,000", "$18,000", "$50,000", "约为可控成本的风险缓冲"],
        ["年度总费用", "$100,000", "$340,000", "$920,000", "基准情景"],
    ]
    add_table(doc, ["费用项目", "第 1 年", "第 2 年", "第 3 年", "说明"], budget_rows, [2450, 1300, 1300, 1300, 2770], 8.6)
    h2(doc, "三年收入预测")
    revenue_rows = [
        ["本地服务完成订单", "$18,000", "$120,000", "$360,000", "成交后推荐/营销服务收入"],
        ["物业与机构项目", "$24,000", "$180,000", "$520,000", "批量再利用、现场运营与报告"],
        ["商家/物业订阅工具", "$0", "$45,000", "$220,000", "稳定交付价值后上线"],
        ["高级便利功能", "$0", "$15,000", "$80,000", "核心互助保持免费"],
        ["年度总收入", "$42,000", "$360,000", "$1,180,000", "不计数据许可收入"],
        ["经营结果", "-$58,000", "$20,000", "$260,000", "税前、融资前基准结果"],
    ]
    add_table(doc, ["收入项目", "第 1 年", "第 2 年", "第 3 年", "说明"], revenue_rows, [2450, 1300, 1300, 1300, 2770], 8.6)
    h2(doc, "资金纪律与阶段门槛")
    add_table(doc, ["阶段", "建议现金上限", "必须证明后才能进入下一阶段"], [
        ["前 30 天", "$5,000–$15,000", "15 次真实交接、5 个服务需求、2 笔真实订单或明确调整结论"],
        ["前 90 天", "$25,000–$45,000", "100 次真实交接、可管理的失约与人工成本、至少 1 个机构试点"],
        ["第 1 年", "$100,000 基准上限", "出现可重复收入，证明一个城市单元有接近可持续的路径"],
        ["第 2 年扩张", "按季度批准", "第二城市在相同安全和成本标准下可复制，不以下载量替代运营结果"],
    ], [1700, 2500, 4920])

    h1(doc, "十二、关键指标、风险与停止条件")
    h2(doc, "核心指标体系")
    add_table(doc, ["维度", "核心指标", "用途"], [
        ["北极星", "每月经取货码核销的真实再利用交接次数", "衡量平台是否真正完成现实任务"],
        ["供需", "合格发布到交接转化率、完成时间、局部密度", "决定品类和区域是否成立"],
        ["信任", "失约率、举报率、严重事件率、申诉恢复率", "决定是否可以扩大规模"],
        ["效率", "每次交接人工分钟数、每个服务需求协调分钟数", "决定运营成本能否下降"],
        ["收入", "有效服务需求率、报价率、成交率、贡献毛利", "决定商业模式是否成立"],
        ["留存", "30 天重复发布/领取、机构续约率", "判断价值是否持续"],
    ], [1400, 4100, 3620])
    h2(doc, "主要风险与应对")
    add_table(doc, ["风险", "早期信号", "应对与停止条件"], [
        ["供需密度不足", "帖子多但申请少，或需求多但供给不足", "缩小区域和品类；连续两轮仍无改善则更换试点区域"],
        ["失约和协调成本过高", "失约率 >25%，每次交接人工 >45 分钟", "提醒、候补、申请额度；无改善则暂停扩量"],
        ["安全或隐私事件", "危险品、骚扰、地址泄露、召回品", "立即下架、保全证据、暂停相关流程并完成复盘"],
        ["服务商质量失控", "投诉、临时加价、许可或保险过期", "暂停推荐，复核主体与合同；严重问题终止合作"],
        ["商业收入弱", "服务需求少或推荐费无法覆盖协调成本", "转向机构项目费、提高预审；验证失败则停止该收入线"],
        ["技术与云成本失控", "媒体、短信、日志或 AI 成本异常", "预算报警、硬上限、压缩、生命周期和供应商复核"],
        ["过早扩张", "新城市没有审核、安全与本地规则", "城市逐个开放；不满足门槛不宣称正式运营"],
    ], [2100, 3100, 4020])

    h1(doc, "十三、立即执行清单")
    add_table(doc, ["期限", "必须完成", "交付物/验收"], [
        ["今天", "确认 USPTO 提交状态和申请序列号；保存脱敏收据", "商标台账中有序列号、付款、类别与截止日"],
        ["3 天内", "确定唯一试点微区域和首批 15 家商家访谈名单", "名单、联系状态和访谈时间"],
        ["7 天内", "完成商家核验表、事故台账、审核与举报 SOP 培训", "每项有负责人和模板"],
        ["14 天内", "完成 15 家商家访谈，取得至少 3 家书面平台价", "书面价格、许可/保险状态和费用条款"],
        ["30 天内", "完成至少 15 次真实交接和 2 笔付费服务订单", "取货码记录、回访、成本与收入复盘"],
        ["60 天内", "完成账户、候补/取消/失约、图片隐私和数据生命周期", "端到端测试与权限审计"],
        ["90 天内", "累计 100 次真实交接并做继续/调整/停止决策", "90 天经营复盘和下一阶段预算审批"],
    ], [1300, 4700, 2420])
    add_callout(doc, "现在最重要的动作", "不要等待完整 App 或更多概念设计。用现有网站和运营后台完成第一批真实交接，记录每一个摩擦、事故、人工分钟和服务需求。真实运营数据将决定 HerbWorld Share 下一步应该开发什么、应该花多少钱。", LIGHT_GREEN, GREEN)

    h1(doc, "附录 A：官方费用与合规依据")
    add_p(doc, "以下链接用于核实官方费用与关键合规要求。市场询价类预算仍需取得律师、CPA、保险经纪、报纸和服务商的正式报价。")
    sources = [
        ("USPTO Trademark Fee Information", "https://www.uspto.gov/trademarks/trademark-fee-information", "基础申请 $350/类；1(b) 使用声明 $150/类；延期 $125/类；维护费等。"),
        ("USPTO Trademark Status and Document Retrieval", "https://tsdr.uspto.gov/", "使用申请序列号查询状态、文件和截止日。"),
        ("New York DOS Foreign LLC Application for Authority", "https://dos.ny.gov/application-authority-foreign-limited-liability-companies", "外国 LLC 权限申请与获授权后 120 天公告要求。"),
        ("New York DOS Biennial Statements", "https://dos.ny.gov/biennial-statements-business-corporations-and-limited-liability-companies", "LLC 每两年申报，官方费用 $9。"),
        ("Apple Developer Program Enrollment", "https://developer.apple.com/programs/enroll/", "组织注册要求与 $99/年会员费。"),
        ("CPSC Resale/Thrift Stores Information Center", "https://www.cpsc.gov/Business--Manufacturing/Business-Education/ResaleThrift-Stores-Information-Center", "二手消费品召回与安全责任信息。"),
        ("CPSC Recalls", "https://www.cpsc.gov/Recalls", "物品批准前的召回查询。"),
        ("New York SHIELD Act", "https://ag.ny.gov/resources/organizations/data-breach-reporting/shield-act", "纽约私人信息合理保障要求。"),
        ("NYSDOT Moving Companies", "https://www.dot.ny.gov/divisions/operating/osss/truck/moving", "纽约州内搬家承运人核验。"),
        ("FMCSA Protect Your Move", "https://www.fmcsa.dot.gov/protect-your-move/search-mover", "跨州搬家承运人核验。"),
    ]
    add_table(doc, ["来源", "网址", "用于核实"], sources, [2600, 3550, 2970], 8.0)

    h1(doc, "附录 B：内部资料依据")
    add_bullets(doc, [
        "/Users/kalien/herbworld/HERBWORLD_SHARE_BUSINESS_PLAN_2026.md",
        "/Users/kalien/herbworld/HERBWORLD_PRODUCT_DOCTRINE.md",
        "/Users/kalien/herbworld/HERBWORLD_30_DAY_PILOT_PLAYBOOK.md",
        "/Users/kalien/herbworld/HERBWORLD_COMMUNITY_COMMERCE_RESEARCH_2026-06-10.md",
        "/Users/kalien/herbworld/LEGAL_AND_SAFETY_RESEARCH_2026-06-12.md",
        "/Users/kalien/herbworld/LAUNCH_OPERATIONS_SOP.md",
        "/Users/kalien/herbworld/app、lib、database 与 ios 当前实现。",
    ])
    add_p(doc, "文件结束", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED)

    props = doc.core_properties
    props.title = "HerbWorld Share 项目详细策划书"
    props.subject = "产品、运营、合规、预算与执行计划"
    props.author = "Polaris Global L.L.C."
    props.keywords = "HerbWorld Share, Manhattan, reuse, business plan"
    props.comments = "Generated from current HerbWorld project materials and official fee sources."

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
